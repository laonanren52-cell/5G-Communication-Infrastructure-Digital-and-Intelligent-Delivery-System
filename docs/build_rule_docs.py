from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(document: Document, title: str, subtitle: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle_paragraph = document.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Microsoft YaHei"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(46, 116, 181)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_rules_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Inches(widths[index])
            set_cell_text(cells[index], value)

    document.add_paragraph()


def build_program_check_doc() -> None:
    document = Document()
    style_document(document, "程序检查规则", "轻量智能审查模块用于定位图纸/设计表中的工程风险、字段缺失和资源冲突。")

    add_heading(document, "输出分级")
    add_rules_table(
        document,
        ["等级", "含义", "处理建议"],
        [
            ["错误", "数据不满足正式生成条件，例如缺字段、重复站点、无效距离或资源重复。", "必须修正后重新上传。"],
            ["高风险", "数据可读取，但存在明显施工风险，例如超长路由或临时取电。", "生成前建议复核并补充措施。"],
            ["提醒", "不会直接阻断预览，但需要作为交底关注点。", "可继续预览，正式交付前建议完善。"],
            ["正常", "对应检查项未发现问题。", "保持当前数据口径。"],
        ],
        [1.0, 3.0, 2.3],
    )

    add_heading(document, "审查项规则")
    add_rules_table(
        document,
        ["审查项", "检查字段", "规则", "结果"],
        [
            ["字段完整性审查", "site_id、site_name、site_type、cable_distance_m、power_supply_mode、fiber_core_count、fiber_start、fiber_end", "字段不存在或存在空值。", "错误；完整且无空值为正常。"],
            ["线缆距离异常审查", "cable_distance_m / 线缆敷设距离", "无法转数字或小于等于 0；大于 500；100 到 500。", "分别为错误、高风险、提醒。"],
            ["站点编号重复审查", "site_id / 站点编号", "非空站点编号重复。", "错误；无重复为正常。"],
            ["纤芯占用异常审查", "fiber_core_count、fiber_core_id、fiber_core_no、fiber_start、fiber_end", "纤芯数量无效、小于等于 0、同一路由纤芯编号重复、纤芯数量大于等于 24。", "分别为错误、错误、错误、提醒。"],
            ["取电方式施工风险审查", "power_supply_mode / 取电方式", "临时取电、转供电、未知取值。", "分别为高风险、提醒、提醒。"],
            ["资源配置冲突审查", "device_id、port_id、resource_id、fiber_core_id", "资源字段重复非空值，或未提供资源字段。", "重复为错误；未提供为提醒。"],
            ["批量文件来源追踪", "source_file / 文件来源", "批量上传时给每条审查明细追加来源文件。", "用于定位问题来自哪一张图纸。"],
        ],
        [1.4, 1.7, 2.2, 1.0],
    )

    add_heading(document, "页面筛选")
    for item in ["全部明细：显示所有审查结果。", "错误项：只显示阻断正式生成的问题。", "高风险项：只显示需要重点复核的施工风险。", "提醒项：只显示交底关注点。", "正常项：只显示通过的检查项。"]:
        add_bullet(document, item)

    document.save(OUTPUT_DIR / "程序检查规则.docx")


def build_validation_doc() -> None:
    document = Document()
    style_document(document, "数据校验规则", "基础数据校验用于确认上传表是否具备进入 BOM 估算和施工资料生成的最低结构条件。")

    add_heading(document, "必填字段")
    add_rules_table(
        document,
        ["字段", "规则", "异常处理"],
        [
            ["站点编号", "必须存在。", "缺失时报错，要求修正表头或补充字段。"],
            ["站点类型", "必须存在。", "缺失时报错，要求修正表头或补充字段。"],
            ["AAU型号", "必须存在。", "缺失时报错，要求修正表头或补充字段。"],
            ["BBU型号", "必须存在。", "缺失时报错，要求修正表头或补充字段。"],
            ["线缆敷设距离", "必须存在。", "缺失时报错，无法进入正式生成。"],
            ["取电方式", "必须存在。", "缺失时报错，无法进入正式生成。"],
        ],
        [1.5, 2.2, 2.6],
    )

    add_heading(document, "字段值校验")
    add_rules_table(
        document,
        ["校验对象", "规则", "处理方式"],
        [
            ["线缆敷设距离", "无法转换为数字。", "提示警告，并按 0 米参与估算。"],
            ["线缆敷设距离", "小于 0。", "提示警告，并按 0 米参与估算。"],
            ["线缆敷设距离", "正常数字。", "保留数值并参与 BOM 估算。"],
            ["站点编号", "存在空值。", "提示警告，建议补齐后再用于正式交付。"],
        ],
        [1.5, 2.3, 2.5],
    )

    add_heading(document, "校验与审查的关系")
    for item in [
        "数据校验是基础结构校验，主要检查中文必填字段和核心数值可用性。",
        "轻量智能审查是工程风险审查，会进一步检查标准字段、距离区间、纤芯、取电和资源冲突。",
        "校验错误或审查错误会阻止正式施工资料生成，但页面仍可保留预览提示。",
    ]:
        add_bullet(document, item)

    document.save(OUTPUT_DIR / "数据校验规则.docx")


if __name__ == "__main__":
    build_program_check_doc()
    build_validation_doc()
